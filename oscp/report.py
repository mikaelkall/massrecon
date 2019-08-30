#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
  Tries to automate the OSCP report writing from CherryTree data.
"""
__author__ = 'kall.micke@gmail.com'

import time
import os
import re
import sys
import re
from docx import Document

from shutil import copyfile

try:
    from librecon.configuration import *
    from librecon.utils import *
except:
    from configuration import *
    from utils import *

class Report:

    def __init__(self):

        home = str(Path.home())
        self.report_dir = '%s/.massrecon/' % home
        self.cfg = Configuration()
        self.oscp_email = self.cfg.config.get('oscp', 'email')
        self.oscp_osid = self.cfg.config.get('oscp', 'osid')
        self.report_file = "%s/OSCP-%s-Exam-Report.docx" % (self.report_dir, self.oscp_osid)

    def __docx_replace_regex(self, doc_obj, regex, replace):

        for p in doc_obj.paragraphs:
            if regex.search(p.text):
                inline = p.runs
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        text = regex.sub(replace, inline[i].text)
                        inline[i].text = text

        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.__docx_replace_regex(cell, regex, replace)

    def replace(self, search, replace):

        regex1 = re.compile(search)
        doc = Document(self.report_file)
        self.__docx_replace_regex(doc, regex1, replace)
        doc.save(self.report_file)

    def create(self):

        if os.path.isfile('template/OSCP-OS-XXXXX-Exam-Report_Template3.2.docx') is False:
            utils().puts('error', 'Template file not found')
            sys.exit(0)

        copyfile('template/OSCP-OS-XXXXX-Exam-Report_Template3.2.docx', self.report_file)
        utils().puts('success', "Created: %s" % self.report_file)

        self.replace(r"example@example.example", self.oscp_email)
        self.replace(r'OSID: XXXXX', self.oscp_osid)

